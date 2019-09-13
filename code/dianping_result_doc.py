from docx import Document
import sys

from docx import Document
from docx.shared import Inches
#from imp import reload
from DB_Handle import get_AllShops_specified_shop_and_city,get_Chain_shops_for_specified_food
#num_list = [1.5, 0.6, 7.8, 6]
#plt.bar(range(len(num_list)), num_list)
#plt.show()
from excel_handle import get_fastest_growth_chain_shop_info
from dianping_matplotlib_picture import draw_statitics_chain_shops_pic
def main():
    #reload(sys)
    #sys.setdefaultencoding('utf-8')

    # 创建文档对象
    document = Document()

    # 设置文档标题，中文要用unicode字符串
    document.add_heading(u'我的一个新文档', 0)

    # 往文档中添加段落
    p = document.add_paragraph('This is a paragraph having some ')
    p.add_run('bold ').bold = True
    p.add_run('and some ')
    p.add_run('italic.').italic = True

    # 添加一级标题
    document.add_heading(u'一级标题, level = 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    # 添加无序列表
    document.add_paragraph('first item in unordered list', style='ListBullet')

    # 添加有序列表
    document.add_paragraph('first item in ordered list', style='ListNumber')
    document.add_paragraph('second item in ordered list', style='ListNumber')
    document.add_paragraph('third item in ordered list', style='ListNumber')

    # 添加图片，并指定宽度
    document.add_picture('test2.jpg', width=Inches(7))

    # 添加表格: 1行3列
    table = document.add_table(rows=1, cols=3)
    # 获取第一行的单元格列表对象
    hdr_cells = table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'Tel'
    # 为表格添加一行
    new_cells = table.add_row().cells
    new_cells[0].text = 'Tom'
    new_cells[1].text = '19'
    new_cells[2].text = '12345678'

    # 添加分页符
    document.add_page_break()

    # 往新的一页中添加段落
    p = document.add_paragraph('This is a paragraph in new page.')

    # 保存文档
    document.save('demo1.docx')

#def insert_title_and_picture(tile,pic):

if __name__ == '__main__':
    city = '成都市'
    specified_food = '火锅'
    date_duration = 6
    rank_num = 5
    fastest_growth_list = get_fastest_growth_chain_shop_info(specified_food,city,rank_num, date_duration)
    title = str(date_duration) + '天内' + city + specified_food + '前' + str(rank_num) + '名'
    draw_statitics_chain_shops_pic(title, fastest_growth_list)

    #获得每个店的信息
    for each_info in fastest_growth_list:
        shop_name = each_info[0]
        target_list = get_AllShops_specified_shop_and_city(shop_name,city)
        for each_shop in target_list:
            print(each_shop)
        print('*' * 100)